import io
import re
import os
import json
from pypdf import PdfReader
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai

# A collection of common technical and professional skills for extraction
COMMON_SKILLS = [
    # Programming Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "go", "rust", "php", "swift", "kotlin", "scala", "html", "css", "sql", "r",
    # Frameworks & Libraries
    "react", "angular", "vue", "next.js", "node.js", "express", "django", "flask", "fastapi", "spring boot", "laravel", "rails", "asp.net", "jquery",
    "bootstrap", "tailwind", "pytorch", "tensorflow", "keras", "scikit-learn", "pandas", "numpy", "opencv", "redux", "graphql",
    # Databases & Cloud
    "mongodb", "postgresql", "mysql", "sqlite", "redis", "elasticsearch", "cassandra", "mariadb", "firebase", "aws", "azure", "gcp", "google cloud",
    "docker", "kubernetes", "jenkins", "git", "github", "gitlab", "terraform", "ansible", "linux", "unix", "nginx", "apache",
    # Methodologies & General
    "agile", "scrum", "kanban", "ci/cd", "devops", "rest api", "microservices", "system design", "data structures", "algorithms",
    "machine learning", "deep learning", "artificial intelligence", "data science", "nlp", "computer vision", "ui/ux", "product management",
    "project management", "business analysis", "qa testing", "unit testing", "cybersecurity", "blockchain"
]

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    
    if ext == ".pdf":
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text_list = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_list.append(t)
        text = "\n".join(text_list)
        
    elif ext == ".docx":
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_list = []
        for para in doc.paragraphs:
            if para.text:
                text_list.append(para.text)
        # Table text extraction
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_list.append(cell.text)
        text = "\n".join(text_list)
        
    elif ext in [".txt", ".md"]:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="ignore")
            
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return text.strip()

def extract_contact_info(text: str):
    """Extract contact information (email, phone, links, and potential name) using Regex."""
    # Find email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    email = email_match.group(0) if email_match else "Not found"
    
    # Find phone (supports formats like +1-234-567-8900, (123) 456-7890, etc.)
    phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    phone = phone_match.group(0) if phone_match else "Not found"
    
    # Find Links (LinkedIn, GitHub)
    links = []
    linkedin_match = re.findall(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\.-]+', text, re.IGNORECASE)
    github_match = re.findall(r'(?:https?://)?(?:www\.)?github\.com/[\w\.-]+', text, re.IGNORECASE)
    
    if linkedin_match:
        links.extend(linkedin_match)
    if github_match:
        links.extend(github_match)
        
    # Attempt name extraction (usually first non-empty line of the text, cleaning it)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    candidate_name = "Candidate Name"
    for line in lines[:5]:
        # Filter out lines that look like links, emails, titles, or headers
        if (len(line) < 40 and 
            "@" not in line and 
            "phone" not in line.lower() and 
            "resume" not in line.lower() and 
            "curriculum" not in line.lower() and 
            "page" not in line.lower() and
            not line.startswith("http")):
            candidate_name = line
            break
            
    return {
        "name": candidate_name,
        "email": email,
        "phone": phone,
        "links": list(set(links))
    }

def analyze_skills_local(text: str, job_desc: str):
    """Analyze matching and missing skills based on the COMMON_SKILLS dictionary."""
    text_lower = text.lower()
    job_desc_lower = job_desc.lower()
    
    jd_skills = []
    matched_skills = []
    missing_skills = []
    
    # Identify skills mentioned in the Job Description
    for skill in COMMON_SKILLS:
        pattern = rf"\b{re.escape(skill)}\b"
        if re.search(pattern, job_desc_lower):
            jd_skills.append(skill)
            
    # If no standard skills match, parse words from JD as potential skills
    if not jd_skills:
        words = re.findall(r'\b[a-zA-Z]{3,15}\b', job_desc_lower)
        stopwords = {"and", "the", "for", "with", "this", "that", "from", "will", "your", "with", "their", "about"}
        candidates = set(words) - stopwords
        jd_skills = list(candidates)[:15] # take top 15
        
    # Find which skills exist in the resume
    for skill in jd_skills:
        pattern = rf"\b{re.escape(skill)}\b"
        if re.search(pattern, text_lower):
            matched_skills.append(skill.title())
        else:
            missing_skills.append(skill.title())
            
    return matched_skills, missing_skills

def local_screener(text: str, job_desc: str) -> dict:
    """Analyze resume locally using TF-IDF and Rule-Based Extraction."""
    contact_info = extract_contact_info(text)
    matched_skills, missing_skills = analyze_skills_local(text, job_desc)
    
    # Calculate similarity score using TF-IDF
    tfidf = TfidfVectorizer(stop_words='english')
    try:
        tfidf_matrix = tfidf.fit_transform([text, job_desc])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Combined score: 40% TF-IDF similarity, 60% skills coverage
        total_skills = len(matched_skills) + len(missing_skills)
        skills_ratio = len(matched_skills) / total_skills if total_skills > 0 else 0.5
        
        raw_score = (similarity * 40) + (skills_ratio * 60)
        match_score = min(100, max(0, int(raw_score)))
    except Exception:
        total_skills = len(matched_skills) + len(missing_skills)
        match_score = int((len(matched_skills) / total_skills * 100)) if total_skills > 0 else 50
        
    strengths = []
    weaknesses = []
    recommendations = []
    
    if len(matched_skills) > 0:
        strengths.append(f"Strong overlap in core skills: {', '.join(matched_skills[:4])}.")
    if match_score >= 70:
        strengths.append("High similarity index with job description vocabulary.")
    else:
        weaknesses.append("Moderate alignment with the overall job vocabulary.")
        
    if len(missing_skills) > 0:
        weaknesses.append(f"Missing critical keywords: {', '.join(missing_skills[:4])}.")
        recommendations.append(f"Incorporate missing skills: {', '.join(missing_skills[:3])} into your experience descriptions.")
    else:
        strengths.append("Covers all identified keyword requirements from the job description.")
        
    recommendations.append("Ensure your resume lists metrics and achievements (e.g. 'improved efficiency by 20%') rather than just tasks.")
    recommendations.append("Tailor your professional summary to align with the core objective of the role.")
    
    questions = [
        f"Can you explain your experience working with {', '.join(matched_skills[:2]) if matched_skills else 'your core technologies'}?",
        "What is the most challenging project you've completed, and what impact did it have?",
        f"How would you approach learning or implementing {missing_skills[0] if missing_skills else 'new concepts required for this role'} in a project?"
    ]
    
    return {
        "candidate_info": contact_info,
        "match_score": match_score,
        "skills_analysis": {
            "matched": matched_skills,
            "missing": missing_skills
        },
        "evaluation": {
            "strengths": strengths,
            "weaknesses": weaknesses,
            "recommendations": recommendations
        },
        "interview_questions": questions,
        "mode": "Quick Local Scan"
    }

def gemini_screener(text: str, job_desc: str, api_key: str) -> dict:
    """Analyze resume using Gemini API for advanced screening and feedback."""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""
You are an expert technical recruiter and resume ATS parsing engine.
Your task is to analyze the provided Resume text and compare it with the Job Description.

Analyze the resume text and the job description, then return a JSON object containing the evaluation.
You must return valid JSON that conforms exactly to this structure:
{{
  "candidate_info": {{
    "name": "Candidate Name or extracted name",
    "email": "Email address or 'Not found'",
    "phone": "Phone number or 'Not found'",
    "links": ["list", "of", "links", "such", "as", "GitHub", "or", "LinkedIn", "or", "portfolio"]
  }},
  "match_score": 75,
  "skills_analysis": {{
    "matched": ["Skill A", "Skill B"],
    "missing": ["Skill C", "Skill D"]
  }},
  "evaluation": {{
    "strengths": ["Strength bullet 1", "Strength bullet 2"],
    "weaknesses": ["Weakness/Gap bullet 1", "Weakness/Gap bullet 2"],
    "recommendations": ["Actionable improvement 1", "Actionable improvement 2"]
  }},
  "interview_questions": [
    "Tailored question 1",
    "Tailored question 2",
    "Tailored question 3"
  ]
}}

Ensure all fields are fully populated and accurate based on the context. Return ONLY the JSON object.

RESUME TEXT:
{text}

JOB DESCRIPTION:
{job_desc}
"""
        response = model.generate_content(
            prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        
        data = json.loads(response.text.strip())
        data["mode"] = "Deep AI Scan"
        return data
        
    except Exception as e:
        print(f"Gemini screening failed: {e}. Falling back to local.")
        fallback = local_screener(text, job_desc)
        fallback["error_info"] = f"Gemini Scan failed: {str(e)}. Reverted to Local Scan."
        return fallback
